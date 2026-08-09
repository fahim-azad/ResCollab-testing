import subprocess
import sys
import re

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    install('python-docx')
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('---'):
            doc.add_page_break()
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], 2)
        elif line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            
            # Basic bold parsing for bullets
            parts = text.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            # Basic bold parsing for regular paragraphs
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(docx_path)
    print(f"Successfully generated {docx_path}")

if __name__ == '__main__':
    markdown_to_docx(r'd:\3-2 class stuff\software lab\project\ResCollab testing\docs\PROJECT_DOCUMENTATION.md', 
                     r'd:\3-2 class stuff\software lab\project\ResCollab testing\docs\PROJECT_DOCUMENTATION.docx')
